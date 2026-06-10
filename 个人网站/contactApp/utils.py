from io import BytesIO

from docx import Document
from docx.shared import Inches


def build_resume_docx(resumes):
    document = Document()
    document.add_heading('个人简历信息表', level=1)

    for resume in resumes:
        document.add_heading('%s - %s' % (resume.name, resume.position), level=2)
        if resume.photo:
            try:
                document.add_picture(resume.photo.path, width=Inches(1.2))
            except Exception:
                document.add_paragraph('证件照片：%s' % resume.photo.name)

        table = document.add_table(rows=6, cols=4)
        table.style = 'Table Grid'
        rows = (
            ('姓名', resume.name, '性别', resume.sex),
            ('身份证号', resume.personID, '出生日期', resume.birth.strftime('%Y-%m-%d')),
            ('邮箱', resume.email, '学历', resume.edu),
            ('毕业学校', resume.school, '专业', resume.major),
            ('申请职位', resume.position, '提交时间', resume.created.strftime('%Y-%m-%d %H:%M:%S')),
            ('学习或工作经历', resume.experience or '', '', ''),
        )
        for row_index, values in enumerate(rows):
            for col_index, value in enumerate(values):
                table.cell(row_index, col_index).text = value
        document.add_paragraph('')

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
